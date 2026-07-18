# Compiles all 17 chapter study guides into one bound volume:
# BADM225-complete-study-guide.docx
# - Master title page
# - Static table of contents (chapter titles + page numbers), filled in from a
#   COM measurement pass — Word's complex TOC/TC field matching proved unreliable
#   across a docxcompose-merged document, so this avoids that fragility entirely.
# - All 17 chapters, each starting on a fresh page
# - One clean volume-wide footer (page X of Y), replacing each chapter's own
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer

BASE = os.path.dirname(__file__)
SITE = os.path.join(BASE, '..', '..')
GREEN = RGBColor(0x0E, 0x4D, 0x33)
INK   = RGBColor(0x21, 0x2A, 0x26)
MUT   = RGBColor(0x5E, 0x6B, 0x64)

CHAPTERS = [
    (1,'Communicating in Today’s Workplace'),(2,'Planning Business Messages'),
    (3,'Organizing and Drafting Business Messages'),(4,'Revising Business Messages'),
    (5,'Short Workplace Messages and Digital Media'),(6,'Positive and Neutral Messages'),
    (7,'Negative Messages'),(8,'Persuasive Messages'),(9,'Informal Reports'),
    (10,'Proposals and Formal Reports'),(11,'Professionalism, Teamwork, and Meetings'),
    (12,'Business Presentations'),(13,'The Job Search and Résumés in the Digital Age'),
    (14,'Interviewing and Follow-Up'),(15,'AI, Agents & Professional Communication'),
    (16,'Leadership Under Pressure & Influence'),(17,'Strategy for Communicators'),
]

def field(paragraph, instr, size=9, color=MUT, bold=False):
    r = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    r._r.append(f1); r._r.append(it); r._r.append(f2)
    r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = 'Calibri'; r.font.bold = bold

def build_master():
    doc = Document()
    st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(11)
    st.font.color.rgb = INK
    st.paragraph_format.line_spacing = 1.15
    st.paragraph_format.space_after = Pt(8)
    for name, size in [('Title', 30), ('Heading 1', 16), ('Heading 2', 13)]:
        s = doc.styles[name]; s.font.name = 'Cambria'; s.font.size = Pt(size)
        s.font.bold = True; s.font.color.rgb = GREEN
    doc.core_properties.title = 'BADM 225 — Professional Communication for Business: Complete Study Guide'
    doc.core_properties.author = 'Derek D. Bussan, Ph.D., M.B.A.'
    doc.core_properties.language = 'en-US'

    k = doc.add_paragraph(); r = k.add_run('BADM 225 · PROFESSIONAL COMMUNICATION FOR BUSINESS')
    r.font.size = Pt(10); r.font.color.rgb = MUT; r.font.name = 'Consolas'
    doc.add_paragraph('Complete Study Guide', style='Title')
    p = doc.add_paragraph(); r = p.add_run('All seventeen chapters, compiled into one volume')
    r.italic = True; r.font.size = Pt(14); r.font.color.rgb = MUT
    p2 = doc.add_paragraph(); r2 = p2.add_run(
        'Original instructional material written by the instructor. Not affiliated with, '
        'and containing no content from, any textbook or publisher. Nothing in this course’s '
        'companion practice site uploads or stores student data — completion codes are the only '
        'record that reaches the instructor.')
    r2.font.size = Pt(10); r2.font.color.rgb = MUT
    p3 = doc.add_paragraph(); r3 = p3.add_run(
        '© 2026 Derek D. Bussan. All rights reserved. Free for individual student and instructor '
        'use; institutional adoption or redistribution requires a license from the author.')
    r3.font.size = Pt(10); r3.font.color.rgb = MUT; r3.font.bold = True
    doc.add_page_break()

    doc.add_paragraph('Table of Contents', style='Heading 1')
    for num, title in CHAPTERS:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.tab_stops.add_tab_stop(Inches(6.4), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(f'Chapter {num}: {title}\t000')
        r.font.size = Pt(11); r.font.color.rgb = INK
        r.font.name = 'Calibri'
    doc.add_page_break()
    return doc

def add_chapter_bookmark(paragraph, num):
    """Wrap the chapter's title paragraph in a named bookmark so a later COM pass can
    look up its actual page number via Bookmarks(name).Range.Information(wdActiveEndPageNumber)."""
    name = f'Chapter{num}'
    start = OxmlElement('w:bookmarkStart'); start.set(qn('w:id'), str(num)); start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd'); end.set(qn('w:id'), str(num))
    ppr = paragraph._p.find(qn('w:pPr'))
    insert_at = 1 if ppr is not None else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)

def set_volume_footer(doc, total_pages=None):
    """PAGE stays a live field (proven reliable across all 17 individual chapter guides).
    NUMPAGES is written as static text once the true total is known via a COM measurement
    pass — recomputing NUMPAGES for a footer on a docxcompose-merged document proved
    unreliable to force via headless automation in this environment, evaluating in an
    isolated context rather than the real page flow. A static total is the correct
    trade-off for a finished, compiled edition."""
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run('© 2026 Derek D. Bussan · BADM 225 Complete Study Guide  —  Page ')
    r1.font.size = Pt(9); r1.font.color.rgb = MUT; r1.font.name = 'Calibri'
    field(p, 'PAGE')
    r2 = p.add_run(' of ' + (str(total_pages) if total_pages else '000'))
    r2.font.size = Pt(9); r2.font.color.rgb = MUT; r2.font.name = 'Calibri'

def main():
    master = build_master()
    composer = Composer(master)
    for i, (num, title) in enumerate(CHAPTERS):
        path = os.path.join(SITE, 'guides', f'ch{num:02d}-study-guide.docx')
        sub = Document(path)
        if i > 0:
            master.add_page_break()
        composer.append(sub)
        prefix = f'Chapter {num}:'
        match = None
        for p in reversed(master.paragraphs):
            if p.style.name == 'Title' and p.text.strip().startswith(prefix):
                match = p; break
        if match is None:
            print(f'  WARNING: could not locate title paragraph for ch{num:02d} to bookmark')
        else:
            add_chapter_bookmark(match, num)
        print(f'appended ch{num:02d}: {title}')
    set_volume_footer(master)
    out = os.path.join(SITE, 'guides', 'BADM225-complete-study-guide.docx')
    composer.save(out)
    print('Saved', out)

if __name__ == '__main__':
    main()
