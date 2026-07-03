# guidelib.py — shared builder for BADM 225 chapter study guides (accessible Word docs)
# Accessibility: real Title/Heading styles (checkers require them), real list styles,
# document language + core properties, high-contrast palette matching the course site.
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

GREEN = RGBColor(0x0E, 0x4D, 0x33)
INK   = RGBColor(0x21, 0x2A, 0x26)
MUT   = RGBColor(0x5E, 0x6B, 0x64)

def new_doc(chapter_num, chapter_title, subtitle):
    doc = Document()
    # base styles
    st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(11)
    st.font.color.rgb = INK
    st.paragraph_format.line_spacing = 1.15
    st.paragraph_format.space_after = Pt(8)
    for name, size in [('Title', 26), ('Heading 1', 16), ('Heading 2', 13)]:
        s = doc.styles[name]; s.font.name = 'Cambria'; s.font.size = Pt(size)
        s.font.bold = True; s.font.color.rgb = GREEN
    doc.styles['Heading 1'].paragraph_format.space_before = Pt(18)
    doc.styles['Heading 1'].paragraph_format.space_after = Pt(8)
    doc.styles['Heading 2'].paragraph_format.space_before = Pt(12)
    # metadata (accessibility checkers look for a document title)
    doc.core_properties.title = f'BADM 225 Study Guide — Chapter {chapter_num}: {chapter_title}'
    doc.core_properties.author = 'Derek D. Bussan, Ph.D., M.B.A.'
    doc.core_properties.language = 'en-US'
    # cover block
    k = doc.add_paragraph(); r = k.add_run('BADM 225 · PROFESSIONAL COMMUNICATION FOR BUSINESS · STUDY GUIDE')
    r.font.size = Pt(9); r.font.color.rgb = MUT; r.font.name = 'Consolas'
    doc.add_paragraph(f'Chapter {chapter_num}: {chapter_title}', style='Title')
    p = doc.add_paragraph(); r = p.add_run(subtitle)
    r.italic = True; r.font.size = Pt(12); r.font.color.rgb = MUT
    p = doc.add_paragraph()
    r = p.add_run('Original instructional material written by the instructor. Not affiliated with, '
                  'and containing no content from, any textbook or publisher.')
    r.font.size = Pt(9); r.font.color.rgb = MUT
    doc.add_page_break()
    _page_numbers(doc, chapter_num)
    return doc

def _field(paragraph, instr, size=9):
    r = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    r._r.append(f1); r._r.append(it); r._r.append(f2)
    r.font.size = Pt(size); r.font.color.rgb = MUT; r.font.name = 'Calibri'

def _page_numbers(doc, chapter_num):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'BADM 225 Study Guide · Chapter {chapter_num}  —  Page ')
    r1.font.size = Pt(9); r1.font.color.rgb = MUT
    _field(p, 'PAGE')
    r2 = p.add_run(' of ')
    r2.font.size = Pt(9); r2.font.color.rgb = MUT
    _field(p, 'NUMPAGES')

def h1(doc, text): doc.add_paragraph(text, style='Heading 1')
def h2(doc, text): doc.add_paragraph(text, style='Heading 2')

def para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead + ' '); r.bold = True; r.font.color.rgb = GREEN
    p.add_run(text)
    return p

def bullets(doc, items):
    for it in items:
        lead, rest = (it if isinstance(it, tuple) else (None, it))
        p = doc.add_paragraph(style='List Bullet')
        if lead:
            r = p.add_run(lead + ' '); r.bold = True; r.font.color.rgb = GREEN
        p.add_run(rest)

def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style='List Number')

def keyterms(doc, terms):
    h1(doc, 'Key terms')
    for term, definition in terms:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(term + ' — '); r.bold = True
        p.add_run(definition)

def references(doc, refs):
    h1(doc, 'References')
    intro = doc.add_paragraph()
    r = intro.add_run('The following published sources inform this chapter and are recommended for deeper study.')
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = MUT
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)  # hanging indent
        p.paragraph_format.space_after = Pt(8)

def figure(doc, img_path, caption, alt):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); pic = run.add_picture(img_path, width=Inches(6.0))
    pic._inline.docPr.set('descr', alt)   # alt text for screen readers / checkers
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = MUT

def quiz(doc, questions, answers):
    h1(doc, 'Self-check quiz')
    p = doc.add_paragraph(); r = p.add_run('Answer before looking at the key on the final page. '
        'If you miss one, reread that section — then take the full practice set on the course site.')
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = MUT
    for i, (q, opts) in enumerate(questions, 1):
        pq = doc.add_paragraph(); rq = pq.add_run(f'{i}. {q}'); rq.bold = True
        for letter, opt in zip('abcd', opts):
            doc.add_paragraph(f'{letter}) {opt}').paragraph_format.left_indent = Inches(0.4)
    doc.add_page_break()
    h2(doc, 'Answer key')
    key = doc.add_paragraph()
    key.add_run('   '.join(f'{i}. {a}' for i, a in enumerate(answers, 1))).bold = True

def finish(doc, path, site_note=True):
    if site_note:
        p = doc.add_paragraph()
        r = p.add_run('Practice what you just read: course site → this chapter → Practice, '
                      'then the graded homework before the weekly deadline.')
        r.bold = True; r.font.color.rgb = GREEN
    doc.save(path)
    print('Saved', path)
